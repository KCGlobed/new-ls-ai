import structlog
from docx import Document

from app.rag.loaders.base import BaseLoader, LoadedPage

logger = structlog.get_logger(__name__)


class DocxLoader(BaseLoader):
    def load(self, file_path: str) -> list[LoadedPage]:
        doc = Document(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        logger.info("docx_loaded", file=file_path, paragraph_count=len(doc.paragraphs))
        return [
            LoadedPage(
                text=full_text,
                page_number=1,
                metadata={"source": file_path},
            )
        ]
